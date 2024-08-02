import tkinter as tk
from tkinter import filedialog, messagebox
import docx
import datetime as dt
import docx2pdf

class InvoiceAutomation:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title('Invoice Automation')
        self.root.geometry('500x600')

        # Labels
        self.partner_label = tk.Label(self.root, text='Partner')
        self.partner_street_label = tk.Label(self.root, text='Partner Street')
        self.partner_zip_city_country_label = tk.Label(self.root, text="Partner ZIP Country")
        self.invoice_number_label = tk.Label(self.root, text='Invoice Number')
        self.service_description_label = tk.Label(self.root, text='Service Description')
        self.service_amount_label = tk.Label(self.root, text='Service Amount')
        self.service_single_price_label = tk.Label(self.root, text='Service Single Price')
        self.payment_method_label = tk.Label(self.root, text='Payment Method')

        # Payment methods
        self.payment_methods = {
            'Main Bank': {
                'Recipient': 'XYS Company',
                'Bank': 'Hello World Bank',
                'IBAN': 'XY12 3456 7890 1234',
                'BIC': 'ABCDEFGH'
            },
            'Second Bank': {
                'Recipient': 'HDFC Company',
                'Bank': 'Bye World Bank',
                'IBAN': 'AB02 3456 7890 1234',
                'BIC': 'ZYWXYZUTS'
            },
            'Third Bank': {
                'Recipient': 'HARI Company',
                'Bank': 'ALOHA World Bank',
                'IBAN': 'MN02 3456 8790 1934',
                'BIC': 'ASDFGHJK'
            }
        }

        # Entry fields
        self.partner_entry = tk.Entry(self.root)
        self.partner_street_entry = tk.Entry(self.root)
        self.partner_zip_city_country_entry = tk.Entry(self.root)
        self.invoice_number_entry = tk.Entry(self.root)
        self.service_description_entry = tk.Entry(self.root)
        self.service_amount_entry = tk.Entry(self.root)
        self.service_single_price_entry = tk.Entry(self.root)

        # Payment method dropdown
        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set('Main Bank')
        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method, *self.payment_methods.keys())

        # Create button
        self.create_button = tk.Button(self.root, text='Create Invoice', command=self.create_invoice)

        # Layout
        padding_options = {'fill': 'x', 'expand': True, 'padx': 5, 'pady': 2}
        self.partner_label.pack(padding_options)
        self.partner_entry.pack(padding_options)
        self.partner_street_label.pack(padding_options)
        self.partner_street_entry.pack(padding_options)
        self.partner_zip_city_country_label.pack(padding_options)
        self.partner_zip_city_country_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.service_amount_label.pack(padding_options)
        self.service_amount_entry.pack(padding_options)
        self.service_single_price_label.pack(padding_options)
        self.service_single_price_entry.pack(padding_options)
        self.payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        self.root.mainloop()

    def replace_text(self, paragraph, old_text, new_text):
        if old_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old_text, new_text)

    def create_invoice(self):
        doc = docx.Document('Template.docx')

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        try:
            # Extract and convert values
            amount = float(self.service_amount_entry.get())
            single_price = float(self.service_single_price_entry.get())
            
            # Calculate full price
            full_price = amount * single_price
            
            # Create replacements dictionary
            replacements = {
                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Partner]": self.partner_entry.get(),
                "[Partner Street]": self.partner_street_entry.get(),
                "[Partner ZIP_City_Country]": self.partner_zip_city_country_entry.get(),
                "[Invoice Number]": self.invoice_number_entry.get(),
                "[Service Description]": self.service_description_entry.get(),
                "[Amount]": f"${amount:.2f}",
                "[Single Price]": f"${single_price:.2f}",
                "[Full Price]": f"${full_price:.2f}",
                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[IBAN]": selected_payment_method['IBAN'],
                "[BIC]": selected_payment_method['BIC']
            }
            
            # Debug: Print replacements
            print("Replacements:", replacements)

        except ValueError as e:
            messagebox.showerror(title="Error", message=f'Invalid amount or price! Error: {e}')
            return

        # Replace text in document
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)

        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF documents', '*.pdf')])

        if save_path:
            doc.save('filled.docx')
            try:
                docx2pdf.convert('filled.docx', save_path)
                messagebox.showinfo('Success', 'Invoice created and saved successfully')
            except Exception as e:
                messagebox.showerror(title="Error", message=f'Error converting to PDF: {e}')

if __name__ == '__main__':
    InvoiceAutomation()
